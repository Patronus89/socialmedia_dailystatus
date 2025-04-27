# report_generator.py
# Word document generation logic here
# report_generator.py
from docx import Document
import os
from datetime import datetime

def generate_report(date_str, summarized_texts, output_folder="./daily_reports/"):
    # Create today's folder
    today_folder = os.path.join(output_folder, date_str)
    os.makedirs(today_folder, exist_ok=True)
    
    # Create Word document
    doc = Document()
    doc.add_heading(f"Myanmar Earthquake Monitoring Report - {date_str}", level=1)

    doc.add_heading("Summary of Social Media and News:", level=2)

    for idx, item in enumerate(summarized_texts, 1):
        doc.add_paragraph(f"{idx}. {item}")

    # Save file
    file_path = os.path.join(today_folder, "report.docx")
    doc.save(file_path)
    
    return file_path

print("Generating Word daily report...")